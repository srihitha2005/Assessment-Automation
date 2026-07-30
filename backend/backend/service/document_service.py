"""Compile assessments to DOCX/PDF and parse DOCX back to structured JSON.

DOCX round-trips losslessly: after the visible content we append a hidden
"[JSON_METADATA]" marker followed by a JSON blob of the exact schema stored
in the database. `parse_docx` reads that blob first; the regex-based reader
is only a fallback for teacher-edited documents that stripped the metadata.
"""
from __future__ import annotations

import json
import logging
import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from reportlab.lib import colors
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    ListFlowable,
    ListItem,
    PageBreak,
)

from config import settings

logger = logging.getLogger(__name__)


METADATA_MARKER = "[JSON_METADATA]"


def _image_path(image: str | None) -> Path | None:
    if not image:
        return None
    candidate = Path(image)
    if candidate.is_absolute() and candidate.exists():
        return candidate
    relative = settings.question_bank_root / candidate
    if relative.exists():
        return relative
    stripped = settings.question_bank_root / candidate.name
    return stripped if stripped.exists() else None


class DocumentService:
    def __init__(self) -> None:
        settings.generated_document_root.mkdir(parents=True, exist_ok=True)

    # -------------------------------------------------------------------- DOCX

    def create_docx(self, assessment: dict, questions: list[dict]) -> Path:
        path = settings.generated_document_root / f"Assessment_{assessment['assessmentId']}.docx"
        document = Document()

        document.add_heading(assessment.get("chapterName") or "Assessment", level=0)
        subtitle = document.add_paragraph()
        subtitle.add_run(
            f"{assessment.get('grade', 'Grade —')}  •  Version {assessment.get('version', 1)}  •  "
            f"{assessment.get('totalMarks', 0)} marks"
        ).italic = True

        document.add_heading("Learning outcomes", level=2)
        for outcome in assessment.get("learningOutcomes", []):
            document.add_paragraph(outcome, style="List Bullet")

        document.add_heading("Questions", level=2)
        for question in questions:
            heading = document.add_paragraph()
            run = heading.add_run(
                f"Q{question['questionNumber']}. {question['question']}  ({question['marks']} marks)"
            )
            run.bold = True
            run.font.size = Pt(12)

            for option in question.get("options", []):
                document.add_paragraph(option, style="List Bullet")

            embedded = _image_path(question.get("image"))
            if embedded:
                try:
                    document.add_picture(str(embedded), width=Inches(3.5))
                except Exception:
                    logger.exception("Failed to embed image %s", embedded)
                    document.add_paragraph(f"[image: {question.get('image')}]")
            elif question.get("image"):
                document.add_paragraph(f"[image: {question['image']}]")

            answer_para = document.add_paragraph()
            answer_para.add_run("Answer: ").bold = True
            answer_para.add_run(question.get("answer", ""))

            meta_para = document.add_paragraph()
            meta_run = meta_para.add_run(
                f"Type: {question.get('questionType')}  •  Difficulty: {question.get('difficulty')}"
                f"  •  Bloom: {question.get('bloomsLevel') or question.get('bloomLevel')}"
                f"  •  Outcomes: {', '.join(question.get('learningOutcomes', []))}"
            )
            meta_run.italic = True
            meta_run.font.size = Pt(9)

            document.add_paragraph()

        # Lossless machine-readable footer.
        document.add_page_break()
        marker_para = document.add_paragraph(METADATA_MARKER)
        marker_para.runs[0].font.size = Pt(1)
        payload = json.dumps({"assessment": assessment, "questions": questions}, ensure_ascii=False)
        payload_para = document.add_paragraph(payload)
        payload_para.runs[0].font.size = Pt(1)

        document.save(path)
        logger.info("Wrote DOCX %s", path.name)
        return path

    def parse_docx(self, file_path: Path) -> dict:
        document = Document(file_path)
        paragraphs = [paragraph.text for paragraph in document.paragraphs]
        text = "\n".join(paragraphs)

        marker_idx = text.rfind(METADATA_MARKER)
        if marker_idx != -1:
            blob = text[marker_idx + len(METADATA_MARKER) :].strip()
            try:
                data = json.loads(blob)
                logger.info("Parsed DOCX %s via metadata footer (lossless).", file_path.name)
                return {
                    "assessment": data.get("assessment"),
                    "questions": data.get("questions", []),
                    "source": "metadata",
                }
            except json.JSONDecodeError:
                logger.warning("Metadata footer present but not valid JSON in %s.", file_path.name)

        return {
            "assessment": None,
            "questions": self._parse_visible(paragraphs),
            "source": "regex",
        }

    def _parse_visible(self, paragraphs: list[str]) -> list[dict]:
        questions: list[dict] = []
        current: dict | None = None
        header = re.compile(r"^Q?(\d+)\.\s*(.+?)\s*\((\d+)\s*marks\)\s*$", re.IGNORECASE)
        for line in paragraphs:
            text = line.strip()
            if not text:
                continue
            match = header.match(text)
            if match:
                if current:
                    questions.append(current)
                current = {
                    "questionNumber": int(match.group(1)),
                    "question": match.group(2),
                    "marks": int(match.group(3)),
                    "options": [],
                    "images": [],
                }
                continue
            if not current:
                continue
            if text.lower().startswith("answer:"):
                current["answer"] = text.split(":", 1)[1].strip()
            elif text.startswith("[image:"):
                current["image"] = text[len("[image:") :].rstrip("] ").strip()
            elif text.lower().startswith("type:"):
                continue  # metadata line — safe to skip
            else:
                current["options"].append(text)
        if current:
            questions.append(current)
        return questions

    # ---------------------------------------------------------------------- PDF

    def create_pdf(self, assessment: dict, questions: list[dict]) -> Path:
        if not settings.enable_pdf_export:
            raise RuntimeError("PDF export is disabled via ENABLE_PDF_EXPORT=false.")
        path = settings.generated_document_root / f"Assessment_{assessment['assessmentId']}.pdf"
        styles = getSampleStyleSheet()
        styles.add(
            ParagraphStyle(
                "SmallMeta",
                parent=styles["BodyText"],
                fontSize=8,
                textColor=colors.grey,
                spaceAfter=8,
            )
        )
        doc = SimpleDocTemplate(
            str(path),
            pagesize=LETTER,
            leftMargin=0.75 * inch,
            rightMargin=0.75 * inch,
            topMargin=0.75 * inch,
            bottomMargin=0.75 * inch,
            title=f"Assessment — {assessment.get('chapterName')}",
        )
        story: list = []

        story.append(Paragraph(assessment.get("chapterName") or "Assessment", styles["Title"]))
        story.append(
            Paragraph(
                f"{assessment.get('grade', 'Grade —')} &nbsp;•&nbsp; "
                f"Version {assessment.get('version', 1)} &nbsp;•&nbsp; "
                f"{assessment.get('totalMarks', 0)} marks",
                styles["Italic"],
            )
        )
        story.append(Spacer(1, 0.15 * inch))

        story.append(Paragraph("Learning outcomes", styles["Heading2"]))
        story.append(
            ListFlowable(
                [ListItem(Paragraph(outcome, styles["BodyText"])) for outcome in assessment.get("learningOutcomes", [])],
                bulletType="bullet",
            )
        )
        story.append(Spacer(1, 0.2 * inch))

        story.append(Paragraph("Questions", styles["Heading2"]))
        for question in questions:
            story.append(
                Paragraph(
                    f"<b>Q{question['questionNumber']}.</b> {question['question']} "
                    f"<font color='#666'>({question['marks']} marks)</font>",
                    styles["BodyText"],
                )
            )
            if question.get("options"):
                story.append(
                    ListFlowable(
                        [ListItem(Paragraph(option, styles["BodyText"])) for option in question["options"]],
                        bulletType="bullet",
                        start="circle",
                    )
                )
            embedded = _image_path(question.get("image"))
            if embedded:
                try:
                    story.append(RLImage(str(embedded), width=3.5 * inch, height=2.2 * inch, kind="proportional"))
                except Exception:
                    logger.exception("PDF: could not embed %s", embedded)
                    story.append(Paragraph(f"[image: {question.get('image')}]", styles["Italic"]))
            elif question.get("image"):
                story.append(Paragraph(f"[image: {question['image']}]", styles["Italic"]))
            story.append(Paragraph(f"<b>Answer:</b> {question.get('answer', '')}", styles["BodyText"]))
            story.append(
                Paragraph(
                    f"Type: {question.get('questionType')} • Difficulty: {question.get('difficulty')} • "
                    f"Bloom: {question.get('bloomsLevel') or question.get('bloomLevel')} • "
                    f"Outcomes: {', '.join(question.get('learningOutcomes', []))}",
                    styles["SmallMeta"],
                )
            )

        story.append(PageBreak())
        doc.build(story)
        logger.info("Wrote PDF %s", path.name)
        return path
