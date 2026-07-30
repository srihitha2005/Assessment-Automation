import json
import logging
import re
from pathlib import Path

from docx import Document

from config import settings

logger = logging.getLogger(__name__)


class DocumentService:
    def create_docx(self, assessment: dict, questions: list[dict]) -> Path:
        """Compile an assessment into an editable DOCX file."""
        settings.generated_document_root.mkdir(parents=True, exist_ok=True)
        path = settings.generated_document_root / f"Assessment_{assessment['assessmentId']}.docx"
        document = Document()
        document.add_heading(f"Assessment: {assessment['chapterName']}", level=0)
        document.add_paragraph(f"Grade: {assessment['grade']}")
        document.add_paragraph(f"Total marks: {assessment['marks']}")
        document.add_paragraph("Learning outcomes:")
        for outcome in assessment["learningOutcomes"]:
            document.add_paragraph(outcome, style="List Bullet")

        for question in questions:
            document.add_heading(f"{question['questionNumber']}. {question['question']} ({question['marks']} marks)", level=2)
            for option in question.get("options", []):
                document.add_paragraph(option, style="List Bullet")
            if question.get("image"):
                document.add_paragraph(f"Image: {question['image']}")
            document.add_paragraph(f"Answer: {question['answer']}")
        document.save(path)
        logger.info("Created assessment document %s.", path.name)
        return path

    def parse_docx(self, file_path: Path) -> list[dict]:
        """Extract the simple format written by create_docx for portal ingestion."""
        document = Document(file_path)
        questions: list[dict] = []
        current: dict | None = None
        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            match = re.match(r"(\d+)\.\s*(.*?)\s*\((\d+) marks\)$", text)
            if match:
                if current:
                    questions.append(current)
                current = {"questionNumber": int(match.group(1)), "question": match.group(2), "marks": int(match.group(3)), "options": [], "images": []}
            elif current and text.startswith("Answer:"):
                current["answer"] = text.removeprefix("Answer:").strip()
            elif current and text.startswith("Image:"):
                current["image"] = text.removeprefix("Image:").strip()
            elif current and text:
                current["options"].append(text)
        if current:
            questions.append(current)
        logger.info("Parsed %s questions from %s.", len(questions), file_path.name)
        return questions
